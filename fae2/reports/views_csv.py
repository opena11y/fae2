from __future__ import absolute_import

import csv
import json

from django.http import HttpResponse

from fae2.settings import SITE_URL
from reports.models import WebsiteReport
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def get_implementation_status(impl_status):
    if impl_status in ['C', 'AC', 'AC-MC', 'PI', 'PI-MC', 'NI', 'NI-MC', 'MC']:
        if 'MC' in impl_status:
            return impl_status.strip('MC') + 'R'
        else:
            return impl_status
    else:
        return 'na'


def get_result(result_value):
    if result_value == 5:
        return 'Violation'
    elif result_value == 4:
        return 'Warning'
    elif result_value == 3:
        return 'Manual Check'
    elif result_value == 2:
        return 'Passed'
    elif result_value == 1:
        return 'Not Applicable'


def get_element_result(result_value):
    if result_value == 5:
        return 'Violation'
    elif result_value == 4:
        return 'Warning'
    elif result_value == 3:
        return 'Manual Check'
    elif result_value == 2:
        return 'Hidden'
    elif result_value == 1:
        return 'Pass'


def addMetaData(report_obj, writer, path):
    writer.writerow(['Meta Label', 'Meta Value'])

    writer.writerow(['Title', report_obj.title])
    writer.writerow(['URL', report_obj.url])
    writer.writerow(['Ruleset', report_obj.ruleset.title])
    writer.writerow(['Depth', report_obj.depth])
    writer.writerow(['Pages', report_obj.page_count])
    writer.writerow(['Report URL', SITE_URL + path + '/'])

    writer.writerow([])


def addDocMetaData(report_obj, doc, path, report='', view=''):
    doc.add_paragraph('Report Title: ' + report_obj.title)

    p = doc.add_paragraph()
    p.add_run('URL: ')
    add_hyperlink(p, report_obj.url, report_obj.url, '327f97', True)
    p.add_run('     Depth: ' + str(report_obj.depth))
    p.add_run('     Ruleset: ')
    add_hyperlink(p, SITE_URL + path + '/ruleset/' + report_obj.ruleset.slug, report_obj.ruleset.title, '327f97', True)

    p = doc.add_paragraph('Pages: ' + str(report_obj.page_count) + '   Page Limit: ' + str(report_obj.max_pages) + '    Non-HTML Files: ' +
                          str(report_obj.excluded_urls.count()))
    p.add_run('     ')
    add_hyperlink(p, SITE_URL + path + '/summary/' + report + '/' + view + '/urls/', 'URL Evaluation Information', '327f97', True)


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def ReportRulesViewCSV(request, report, view):
    # Create the HttpResponse object with the appropriate CSV header.
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/csv/', '').replace('/', '-').strip('-') + '.csv"'

    writer = csv.writer(response)

    report_obj = WebsiteReport.objects.get(slug=report)

    addMetaData(report_obj, writer, request.path.replace('/csv/', ''))

    writer.writerow(['Rule Group', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'N/A', 'Score', 'Status'])

    page = False

    if report_obj.page_count == 1:
        page = report_obj.get_first_page()
        if view == 'gl':
            groups = page.page_gl_results.all()
        elif view == 'rs':
            groups = page.page_rs_results.all()
        else:
            groups = page.page_rc_results.all()
    else:
        if view == 'gl':
            groups = report_obj.ws_gl_results.all()
        elif view == 'rs':
            groups = report_obj.ws_rs_results.all()
        else:
            groups = report_obj.ws_rc_results.all()

    for g in groups:
        writer.writerow(
            [g.get_title(), g.rules_violation, g.rules_warning, g.rules_manual_check, g.rules_passed, g.rules_na,
             g.implementation_score, get_implementation_status(g.implementation_status)])

    writer.writerow(
        ['All Report Groups', report_obj.rules_violation, report_obj.rules_warning, report_obj.rules_manual_check,
         report_obj.rules_passed, report_obj.rules_na, report_obj.implementation_score,
         get_implementation_status(report_obj.implementation_status)])
    return response


def ReportRulesViewDocx(request, report, view):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/docx/', '').replace('/', '-').strip('-') + '.docx"'

    # Fetching data
    report_obj = WebsiteReport.objects.get(slug=report)
    page = False

    if report_obj.page_count == 1:
        page = report_obj.get_first_page()
        if view == 'gl':
            groups = page.page_gl_results.all()
        elif view == 'rs':
            groups = page.page_rs_results.all()
        else:
            groups = page.page_rc_results.all()
    else:
        if view == 'gl':
            groups = report_obj.ws_gl_results.all()
        elif view == 'rs':
            groups = report_obj.ws_rs_results.all()
        else:
            groups = report_obj.ws_rc_results.all()

    # Document Formation
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = 457200
        section.bottom_margin = 457200
        section.left_margin = 457200
        section.right_margin = 457200

    addDocMetaData(report_obj, document, request.path.replace('/docx/', ''), report, view)

    document.add_heading('Summary', 1)
    document.add_paragraph('')

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    table.rows[0].height = 304800
    table.columns[0].width = 1371600

    header_cells[1].text = 'Violations'
    header_cells[2].text = 'Warnings'
    header_cells[3].text = 'Manual Checks'
    header_cells[4].text = 'Passed'

    row_cells = table.add_row().cells
    table.rows[1].height = 304800
    row_cells[0].text = 'Number of Rules'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)

    for i in range(5):
        table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(1, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.rows[1].cells[1]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFAEAE"/>'.format(nsdecls('w'))))
    table.rows[1].cells[2]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFEC94"/>'.format(nsdecls('w'))))
    table.rows[1].cells[3]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B4D8E7"/>'.format(nsdecls('w'))))
    table.rows[1].cells[4]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B0E57C"/>'.format(nsdecls('w'))))

    document.add_paragraph('')

    # headers = ['Rule Group', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'N/A', 'Score', 'Status']
    headers = ['Rule Group', 'V', 'W', 'MC', 'P', 'N/A', 'Score', 'Status']

    # First Table
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Colorful Grid'
    header_cells = table.rows[0].cells

    # table.columns[0].width = 1828800
    table.columns[0].width = 1371600

    for i in range(1, len(headers)):
        table.columns[i].width = 750000

    table.rows[0].height = 304800

    for i in range(len(headers)):
        header_cells[i].text = headers[i]
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

        for i in range(len(row_cells)):
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    for i in range(len(row_cells)):
        row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(1, len(table.rows)):
        table.rows[i].height = 304800

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Colorful List'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Colorful Shading'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Dark List'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light List'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_warning)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Grid 1'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Grid 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Grid 3'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium List 1'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium List 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.add_paragraph('')

    # Second Table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    for g in groups:
        row_cells = table.add_row().cells
        row_cells[0].text = g.get_title()
        row_cells[1].text = str(g.rules_violation)
        row_cells[2].text = str(g.rules_warning)
        row_cells[3].text = str(g.rules_manual_check)
        row_cells[4].text = str(g.rules_passed)
        row_cells[5].text = str(g.rules_na)
        row_cells[6].text = str(g.implementation_score)
        row_cells[7].text = get_implementation_status(g.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = 'All Report Groups'
    row_cells[1].text = str(report_obj.rules_violation)
    row_cells[2].text = str(report_obj.rules_violation)
    row_cells[3].text = str(report_obj.rules_manual_check)
    row_cells[4].text = str(report_obj.rules_passed)
    row_cells[5].text = str(report_obj.rules_na)
    row_cells[6].text = str(report_obj.implementation_score)
    row_cells[7].text = get_implementation_status(report_obj.implementation_status)

    document.save(response)

    return response


def ReportRulesGroupViewCSV(request, report, view, group):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/csv/', '').replace('/', '-').strip('-') + '.csv"'

    writer = csv.writer(response)
    report_obj = WebsiteReport.objects.get(slug=report)

    addMetaData(report_obj, writer, request.path.replace('/csv/', ''))

    writer.writerow(
        ['ID', 'Rule Summary', 'Result', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'N/A', 'Score', 'Status'])

    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)

    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)

    else:
        group = report_obj.ws_rc_results.get(slug=group)

    for g in group.ws_rule_results.all():
        writer.writerow(
            [g.rule.nls_rule_id, g.get_title(), get_result(g.result_value), g.pages_violation, g.pages_warning,
             g.pages_manual_check, g.pages_passed, g.pages_na, g.implementation_score,
             get_implementation_status(g.implementation_status)])

    return response


def ReportRulesGroupViewDocx(request, report, view, group):
    document = Document()
    document.add_heading('Summary of ' + report + '-' + view + '-' + group, 1)
    document.add_paragraph('\n')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/docx/', '').replace('/', '-').strip('-') + '.docx"'

    report_obj = WebsiteReport.objects.get(slug=report)

    addDocMetaData(report_obj, document, request.path.replace('/docx/', ''))

    headers = ['ID', 'Rule Summary', 'Result', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'N/A', 'Score',
               'Status']
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)

    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)

    else:
        group = report_obj.ws_rc_results.get(slug=group)

    for g in group.ws_rule_results.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(g.rule.nls_rule_id)
        row_cells[1].text = g.get_title()
        row_cells[2].text = get_result(g.result_value)
        row_cells[3].text = str(g.pages_violation)
        row_cells[4].text = str(g.pages_warning)
        row_cells[5].text = str(g.pages_manual_check)
        row_cells[6].text = str(g.pages_passed)
        row_cells[7].text = str(g.pages_na)
        row_cells[8].text = str(g.implementation_score)
        row_cells[9].text = get_implementation_status(g.implementation_status)

    document.save(response)
    return response


def ReportRulesGroupRuleViewCSV(request, report, view, group, rule):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/csv/', '').replace('/', '-').strip('-') + '.csv"'

    writer = csv.writer(response)
    report_obj = WebsiteReport.objects.get(slug=report)

    addMetaData(report_obj, writer, request.path.replace('/csv/', ''))

    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)
    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)
    else:
        group = report_obj.ws_rc_results.get(slug=group)

    ws_rule_result = group.ws_rule_results.get(slug=rule)

    writer.writerow(
        ['Page', 'Page Title', 'Result', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'Score', 'Status'])

    for wsr in ws_rule_result.page_rule_results.all():
        writer.writerow(
            [wsr.page_result.page_number, wsr.page_result.title, get_result(wsr.result_value), wsr.elements_violation,
             wsr.elements_warning, wsr.elements_mc_identified, wsr.elements_passed, wsr.implementation_score,
             get_implementation_status(wsr.implementation_status)])

    writer.writerow([None, 'All Pages', get_result(ws_rule_result.result_value), ws_rule_result.elements_violation,
                     ws_rule_result.elements_warning, ws_rule_result.elements_mc_identified,
                     ws_rule_result.elements_passed, ws_rule_result.implementation_score,
                     get_implementation_status(ws_rule_result.implementation_status)])

    return response


def ReportRulesGroupRuleViewDocx(request, report, view, group, rule):
    document = Document()
    document.add_heading('Summary of ' + report + '-' + view + '-' + group + '-' + rule, 1)
    document.add_paragraph('\n')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/docx/', '').replace('/', '-').strip('-') + '.docx"'

    report_obj = WebsiteReport.objects.get(slug=report)

    addDocMetaData(report_obj, document, request.path.replace('/docx/', ''))

    headers = ['Page', 'Page Title', 'Result', 'Violations', 'Warnings', 'Manual Check', 'Passed', 'Score', 'Status']
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)
    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)
    else:
        group = report_obj.ws_rc_results.get(slug=group)

    ws_rule_result = group.ws_rule_results.get(slug=rule)

    for wsr in ws_rule_result.page_rule_results.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(wsr.page_result.page_number)
        row_cells[1].text = wsr.page_result.title
        row_cells[2].text = get_result(wsr.result_value)
        row_cells[3].text = str(wsr.elements_violation)
        row_cells[4].text = str(wsr.elements_warning)
        row_cells[5].text = str(wsr.elements_mc_identified)
        row_cells[6].text = str(wsr.elements_passed)
        row_cells[7].text = str(wsr.implementation_score)
        row_cells[8].text = get_implementation_status(wsr.implementation_status)

    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = 'All Pages'
    row_cells[2].text = get_result(ws_rule_result.result_value)
    row_cells[3].text = str(ws_rule_result.elements_violation)
    row_cells[4].text = str(ws_rule_result.elements_warning)
    row_cells[5].text = str(ws_rule_result.elements_mc_identified)
    row_cells[6].text = str(ws_rule_result.elements_passed)
    row_cells[7].text = str(ws_rule_result.implementation_score)
    row_cells[8].text = get_implementation_status(ws_rule_result.implementation_status)

    document.save(response)
    return response


def ReportRulesGroupRulePageViewCSV(request, report, view, group, rule, page):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/csv/', '').replace('/', '-').strip('-') + '.csv"'

    writer = csv.writer(response)
    report_obj = WebsiteReport.objects.get(slug=report)

    addMetaData(report_obj, writer, request.path.replace('/csv/', ''))
    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)
    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)
    else:
        group = report_obj.ws_rc_results.get(slug=group)

    ws_rule_result = group.ws_rule_results.get(slug=rule)

    page_rule_result = ws_rule_result.page_rule_results.get(page_result__page_number=page)

    writer.writerow(['Element Identifier', 'Result', 'Element Position', 'Message'])

    for prr in json.loads(page_rule_result.element_results_json):
        writer.writerow(
            [prr['element_identifier'], get_element_result(int(prr['result_value'])), prr['ordinal_position'],
             prr['message']])

    return response


def ReportRulesGroupRulePageViewDocx(request, report, view, group, rule, page):
    document = Document()
    document.add_heading('Summary of ' + report + '-' + view + '-' + group + '-' + rule, 1)
    document.add_paragraph('\n')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="' + \
                                      request.path.replace('/docx/', '').replace('/', '-').strip('-') + '.docx"'

    report_obj = WebsiteReport.objects.get(slug=report)

    addDocMetaData(report_obj, document, request.path.replace('/docx/', ''))

    headers = ['Element Identifier', 'Result', 'Element Position', 'Message']
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 2'
    header_cells = table.rows[0].cells

    for i in range(len(headers)):
        header_cells[i].text = headers[i]

    if view == 'gl':
        group = report_obj.ws_gl_results.get(slug=group)
    elif view == 'rs':
        group = report_obj.ws_rs_results.get(slug=group)
    else:
        group = report_obj.ws_rc_results.get(slug=group)

    ws_rule_result = group.ws_rule_results.get(slug=rule)

    page_rule_result = ws_rule_result.page_rule_results.get(page_result__page_number=page)

    for prr in json.loads(page_rule_result.element_results_json):
        row_cells = table.add_row().cells
        row_cells[0].text = prr['element_identifier']
        row_cells[1].text = get_element_result(int(prr['result_value']))
        row_cells[2].text = str(prr['ordinal_position'])
        row_cells[3].text = prr['message']

    document.save(response)

    return response
